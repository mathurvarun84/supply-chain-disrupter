"""
utils.py — Monolithic RAG corpus builder for Capstone Project 8.

This module builds a single ChromaDB collection (`electronics_supply_chain_knowledge`)
from structured Excel data, mitigation playbooks, and static PDF/DOCX reports.

INGESTION (bi-encoder)
----------------------
Each text chunk is embedded with a sentence-transformers *bi-encoder* (fine-tuned
`mathurvarun84/supply-chain-embeddings` from Hugging Face by default; override via
`EMBEDDING_MODEL_PATH`).
ChromaDB's SentenceTransformerEmbeddingFunction encodes query and document texts
*independently* into 384-dim vectors. During `collection.upsert()`, ChromaDB calls
the embedding function internally — callers pass raw text, not vectors.

RETRIEVAL vs RERANKING
----------------------
`query_chroma_rag()` performs Stage-1 bi-encoder retrieval only (cosine distance
over the HNSW index). Stage-2 cross-encoder reranking lives in `src/rag/retriever.py`
(`retrieve_and_rerank`, `rerank_results`) and is used by LangGraph agents for the
named collections in `src/rag/collections.py`.

See also: `src/rag/collections.py` for domain-specific collections (historical
precedents, export control, India sourcing).
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import threading
from pathlib import Path
from typing import Any, Dict, List

import chromadb
import pandas as pd
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.utils.etl_loader import EXCEL_SOURCE, read_excel_sheets
from src.utils.hf_utils import configure_hf_hub_ssl

PROJECT_ROOT = Path(__file__).resolve().parents[2]
CHROMA_DIR = PROJECT_ROOT / "outputs" / "chromadb"
DEFAULT_COLLECTION_NAME = "electronics_supply_chain_knowledge"
PLAYBOOKS_DIR = PROJECT_ROOT / "config" / "playbooks"
STATIC_CONTEXT_DIR = PROJECT_ROOT / "data" / "raw" / "RAG_data"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
DEFAULT_EMBEDDING_REPO = "mathurvarun84/supply-chain-embeddings"
CHUNK_SIZE = 1000
logger = logging.getLogger(__name__)
CHUNK_OVERLAP = 200

_EMBEDDING_WEIGHT_FILES = ("model.safetensors", "pytorch_model.bin")


def _embedding_weights_present(model_dir: Path) -> bool:
    """
    True when a local directory is a loadable sentence-transformers export.

    Weight files alone are insufficient — ST needs modules.json (or a HF
    config.json) so the loader does not mis-detect the path and fail with
    errors like ``No module named 'sentence_transformers.base'``.
    """
    if not model_dir.is_dir():
        return False
    has_weights = any((model_dir / name).exists() for name in _EMBEDDING_WEIGHT_FILES)
    if not has_weights:
        return False
    if (model_dir / "modules.json").is_file():
        return True
    if (model_dir / "config_sentence_transformers.json").is_file():
        return True
    if (model_dir / "config.json").is_file():
        return True
    return False

# Module-level singleton — ChromaDB PersistentClient holds an exclusive file lock
# on chroma.sqlite3. Creating multiple instances in the same process causes WinError 32
# ("file is being used by another process"). A single shared client avoids this.
_chroma_client: chromadb.PersistentClient | None = None
_chroma_lock = threading.Lock()


def get_chroma_client() -> chromadb.PersistentClient:
    """
    Return a process-wide singleton ChromaDB PersistentClient.

    ChromaDB locks `chroma.sqlite3` exclusively; reusing one client avoids
    WinError 32 on Windows when multiple callers open the same store.
    """
    global _chroma_client
    with _chroma_lock:
        if _chroma_client is None:
            CHROMA_DIR.mkdir(parents=True, exist_ok=True)
            _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _chroma_client


def reset_chroma_client() -> None:
    """Drop the cached client (e.g. after --flush rebuild)."""
    global _chroma_client
    with _chroma_lock:
        _chroma_client = None


def resolve_embedding_model_name() -> str:
    """
    Resolve the Hugging Face repo id for the sentence-transformers bi-encoder.

    Priority:
      1. EMBEDDING_MODEL_PATH env var (Hugging Face repo id)
      2. DEFAULT_EMBEDDING_REPO (project fine-tuned model on Hugging Face Hub)
    """
    custom = os.getenv("EMBEDDING_MODEL_PATH", "").strip()
    if custom:
        logger.info("Using embedding model from Hugging Face: %s", custom)
        return custom
    logger.info("Using default fine-tuned embedding model: %s", DEFAULT_EMBEDDING_REPO)
    return DEFAULT_EMBEDDING_REPO


# Process-wide cache, keyed by resolved model_name — constructing
# SentenceTransformerEmbeddingFunction for our fine-tuned Hugging Face repo
# takes ~35-40s the first time (a Hub round-trip to resolve the repo
# revision, not the local weight load, which finishes in under a second).
# query_chroma_rag() previously called get_embedding_model() fresh on every
# single query (3x per L2 news-agent call alone), so whichever agent ran the
# first RAG lookup in a process silently ate that cold-start cost and looked
# slow. Keyed by name (not a single global) so EMBEDDING_MODEL_PATH changing
# mid-process still resolves to a distinct, freshly-built instance.
_embedding_model_cache: Dict[str, SentenceTransformerEmbeddingFunction] = {}
_embedding_model_lock = threading.Lock()


def reset_embedding_model_cache() -> None:
    """Drop cached embedding function instances.

    Call this after pushing new weights to the *same* Hugging Face repo id
    (a same-name update the name-keyed cache above can't detect on its own)
    — e.g. from the admin "Rebuild RAG corpus" action, which is the moment a
    stale in-memory model would otherwise keep re-embedding with old weights
    for the rest of the process's life. A plain env var change to
    EMBEDDING_MODEL_PATH does not need this: it resolves to a new cache key
    automatically.
    """
    with _embedding_model_lock:
        _embedding_model_cache.clear()


def get_embedding_model() -> SentenceTransformerEmbeddingFunction:
    """
    Return the ChromaDB embedding function wrapping the resolved bi-encoder.

    The bi-encoder encodes each chunk and query into a fixed 384-dim vector
    with L2-normalization (`normalize_embeddings=True`) so cosine distance in
    ChromaDB matches semantic similarity. Used at ingest (upsert) and query time.

    Cached per resolved model_name for the life of the process — see
    reset_embedding_model_cache() to force a reload.
    """
    configure_hf_hub_ssl()
    model_name = resolve_embedding_model_name()

    with _embedding_model_lock:
        cached = _embedding_model_cache.get(model_name)
    if cached is not None:
        return cached

    try:
        model = SentenceTransformerEmbeddingFunction(
            model_name=model_name,
            normalize_embeddings=True,
        )
    except Exception as exc:
        if model_name == EMBEDDING_MODEL:
            raise
        logger.warning(
            "Embedding model %s failed to load (%s) — falling back to %s",
            model_name,
            exc,
            EMBEDDING_MODEL,
        )
        model_name = EMBEDDING_MODEL
        with _embedding_model_lock:
            cached = _embedding_model_cache.get(model_name)
        if cached is not None:
            return cached
        model = SentenceTransformerEmbeddingFunction(
            model_name=EMBEDDING_MODEL,
            normalize_embeddings=True,
        )

    with _embedding_model_lock:
        _embedding_model_cache[model_name] = model
    return model


def _chunk_text(text: str) -> List[str]:
    """
    Split long text into overlapping character windows for embedding.

    Prefers breaks at paragraph, newline, or sentence boundaries within the
    second half of each window so chunks stay near CHUNK_SIZE without cutting
    mid-sentence when possible. Overlap (CHUNK_OVERLAP) keeps context that
    straddles a boundary visible in at least one chunk.
    """
    text = re.sub(r"\r\n?", "\n", text).strip()
    if not text:
        return []
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + CHUNK_SIZE)
        if end < len(text):
            split = max(
                text.rfind("\n\n", start + CHUNK_SIZE // 2, end),
                text.rfind("\n", start + CHUNK_SIZE // 2, end),
                text.rfind(". ", start + CHUNK_SIZE // 2, end),
            )
            if split > start:
                end = split + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - CHUNK_OVERLAP)
    return chunks


def _doc_id(source: str, doc_type: str, key: str, chunk_index: int) -> str:
    """Build a stable SHA-256 ID so upserts overwrite the same chunk on rebuild."""
    value = f"{source}|{doc_type}|{key}|{chunk_index}".encode("utf-8")
    return hashlib.sha256(value).hexdigest()


def _add_document(
    documents: List[str],
    metadatas: List[Dict[str, Any]],
    ids: List[str],
    *,
    text: str,
    source: str,
    doc_type: str,
    key: str,
    metadata: Dict[str, Any] | None = None,
) -> None:
    """
    Chunk `text` and append (document, metadata, id) triples to the batch lists.

    Each chunk gets standard metadata (source, type, domain) plus optional
    caller fields. Actual bi-encoder embedding happens later in upsert().
    """
    for index, chunk in enumerate(_chunk_text(text)):
        documents.append(chunk)
        metadatas.append(
            {
                "source": source,
                "type": doc_type,
                "domain": "electronics_semiconductors",
                "dataset_owner": "Varun",
                "beauty_products_included": False,
                "chunk_index": index,
                **(metadata or {}),
            }
        )
        ids.append(_doc_id(source, doc_type, key, index))


def _workbook_documents(
    excel_path: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], Dict[str, int]]:
    """
    Extract RAG documents from the Lite Master Excel workbook.

    Produces five doc types: data dictionary rows, legend entries, aggregated
    disruption event profiles, mitigation recommendations, and semiconductor
    signal events. Returns parallel lists ready for Chroma upsert.
    """
    sheets = read_excel_sheets(excel_path)
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    counts = {
        "data_dictionary": 0,
        "workbook_context": 0,
        "event_profiles": 0,
        "mitigation_guidance": 0,
        "semiconductor_events": 0,
    }

    guide = sheets.get("Column Guide (Lite)")
    if guide is not None:
        guide = guide.dropna(how="all")
        for row_index, row in enumerate(guide.to_dict(orient="records")):
            column = row.get("Column")
            agent = row.get("Agent")
            source_type = row.get("Source Type")
            purpose = row.get("Purpose")

            # v3.1 appendix rows (Known_Disruption_Event, Known_Severity) use Unnamed cols
            if pd.isna(column) or str(column).strip().lower() in ("", "nan", "none"):
                column = row.get("Unnamed: 5")
                if pd.isna(agent) or str(agent).strip().lower() in ("", "nan", "none"):
                    agent = row.get("Unnamed: 4")
                if pd.isna(source_type) or str(source_type).strip().lower() in ("", "nan", "none"):
                    source_type = row.get("Unnamed: 6")
                if pd.isna(purpose) or str(purpose).strip().lower() in ("", "nan", "none"):
                    purpose = row.get("Unnamed: 7")

            if pd.isna(column) or str(column).strip().lower() in ("", "nan", "none"):
                continue

            text = (
                f"Agent: {agent}\n"
                f"Field: {column}\n"
                f"Source type: {source_type}\n"
                f"Purpose: {purpose}"
            )
            _add_document(
                documents,
                metadatas,
                ids,
                text=text,
                source=excel_path.name,
                doc_type="data_dictionary",
                key=f"{row_index}|{column}",
                metadata={"agent": str(agent), "field": str(column)},
            )
            counts["data_dictionary"] += 1

    legend = sheets.get("Legend")
    if legend is not None:
        legend = legend.dropna(how="all")
        for index, row in legend.iterrows():
            text = f"{row.iloc[0]}: {row.iloc[1]}"
            _add_document(
                documents,
                metadatas,
                ids,
                text=text,
                source=excel_path.name,
                doc_type="workbook_context",
                key=str(index),
            )
            counts["workbook_context"] += 1

    master = sheets["Lite Master"]
    grouped = master.groupby("Disruption_Event_Label", dropna=False)
    for label, frame in grouped:
        label_text = "UNLABELLED" if pd.isna(label) else str(label)
        order_dates = pd.to_datetime(frame["Order_Date"], errors="coerce")
        top_regions = ", ".join(
            f"{name} ({count})"
            for name, count in frame["Order_Region"].value_counts().head(5).items()
        )
        top_products = ", ".join(
            f"{name} ({count})"
            for name, count in frame["Product_Name"].value_counts().head(5).items()
        )
        text = (
            f"Electronics disruption profile: {label_text}\n"
            f"Order observations: {len(frame):,}\n"
            f"Date range: {order_dates.min():%Y-%m-%d} to "
            f"{order_dates.max():%Y-%m-%d}\n"
            f"Top affected regions: {top_regions}\n"
            f"Top products: {top_products}\n"
            f"Average composite risk: {frame['Risk_Score_Composite'].mean():.3f}\n"
            f"Average supply disruption index: "
            f"{frame['Supply_Disruption_Index'].mean():.3f}\n"
            f"Average lead-time variance: "
            f"{frame['Lead_Time_Variance_Days'].mean():.2f} days\n"
            f"Average stockout probability: "
            f"{frame['Stockout_Probability_Pct'].mean():.2f}%\n"
            f"Alternate supplier availability: "
            f"{frame['Alternate_Supplier_Available'].mean() * 100:.1f}%"
        )
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="event_profile",
            key=label_text,
            metadata={"event_label": label_text, "observations": len(frame)},
        )
        counts["event_profiles"] += 1

    recommendations = (
        master[
            ["Disruption_Event_Label", "Mitigation_Recommendation"]
        ]
        .dropna()
        .drop_duplicates()
    )
    for _, row in recommendations.iterrows():
        label = str(row["Disruption_Event_Label"])
        recommendation = str(row["Mitigation_Recommendation"])
        text = (
            f"Electronics mitigation guidance\n"
            f"Risk label: {label}\n"
            f"Recommended response: {recommendation}"
        )
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=excel_path.name,
            doc_type="mitigation_guidance",
            key=f"{label}|{recommendation}",
            metadata={"event_label": label},
        )
        counts["mitigation_guidance"] += 1

    signals = sheets.get("Semiconductor Signals")
    if signals is not None:
        signals = signals.dropna(how="all")
        event_rows = signals[
            signals["Known Disruption Event"].notna()
            & ~signals["Known Disruption Event"].astype(str).isin(["—", "-", "None"])
        ]
        event_rows = event_rows.drop_duplicates(
            subset=["Year", "Country", "Company", "Known Disruption Event"]
        )
        for _, row in event_rows.iterrows():
            event = str(row["Known Disruption Event"])
            text = (
                f"Historical semiconductor disruption signal\n"
                f"Year: {row['Year']}\nCountry: {row['Country']}\n"
                f"Company: {row['Company']}\nEvent: {event}\n"
                f"Known severity: {row['Known Severity']}\n"
                f"Supply disruption index: {row['Supply Disruption Index']}\n"
                f"Semiconductor security risk: "
                f"{row['Semiconductor Security Risk']}\n"
                f"Natural disaster risk: {row['Natural Disaster Risk']}\n"
                f"Factory shutdown risk: {row['Factory Shutdown Risk']}\n"
                f"Export control level: {row['Export Control Level']}\n"
                f"Chip price index: {row['Chip Price Index']}"
            )
            key = f"{row['Year']}|{row['Country']}|{row['Company']}|{event}"
            _add_document(
                documents,
                metadatas,
                ids,
                text=text,
                source=excel_path.name,
                doc_type="semiconductor_event",
                key=key,
                metadata={
                    "year": int(row["Year"]),
                    "country": str(row["Country"]),
                    "company": str(row["Company"]),
                    "event": event,
                    "severity": str(row["Known Severity"]),
                },
            )
            counts["semiconductor_events"] += 1

    return documents, metadatas, ids, counts


def _playbook_documents(
    playbooks_dir: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], int]:
    """
    Load mitigation playbook `.txt` files from config/playbooks into chunk lists.

    Each file becomes one or more chunks via `_add_document`. Returns the
    playbook file count for build statistics.
    """
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    count = 0
    if not playbooks_dir.exists():
        return documents, metadatas, ids, count

    for path in sorted(playbooks_dir.glob("*.txt")):
        text = path.read_text(encoding="utf-8")
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=path.name,
            doc_type="mitigation_playbook",
            key=path.stem,
            metadata={"title": path.stem.replace("_", " ")},
        )
        count += 1
    return documents, metadatas, ids, count


def _docx_text(path: Path) -> str:
    """Extract plain text from a DOCX file, including pipe-delimited table rows."""
    document = DocxDocument(path)
    sections: List[str] = []

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    if paragraphs:
        sections.append("\n\n".join(paragraphs))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            sections.append(f"Table {table_index}\n" + "\n".join(rows))

    return "\n\n".join(sections)


def _static_context_documents(
    context_dir: Path,
) -> tuple[List[str], List[Dict[str, Any]], List[str], Dict[str, int]]:
    """
    Ingest static PDF and DOCX reports from data/raw/RAG_data (root level).

    PDFs are chunked per page; DOCX files are chunked as whole documents.
    Skips empty extractions and returns counts by file type.
    """
    documents: List[str] = []
    metadatas: List[Dict[str, Any]] = []
    ids: List[str] = []
    counts = {
        "static_context_files": 0,
        "static_pdf_files": 0,
        "static_pdf_pages": 0,
        "static_docx_files": 0,
    }
    if not context_dir.exists():
        return documents, metadatas, ids, counts

    for path in sorted(context_dir.glob("*.pdf")):
        reader = PdfReader(path)
        indexed_pages = 0
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                continue
            _add_document(
                documents,
                metadatas,
                ids,
                text=text,
                source=path.name,
                doc_type="static_report",
                key=f"{path.stem}|page-{page_number}",
                metadata={
                    "title": path.stem.replace("_", " "),
                    "file_type": "pdf",
                    "page": page_number,
                },
            )
            indexed_pages += 1

        if indexed_pages:
            counts["static_context_files"] += 1
            counts["static_pdf_files"] += 1
            counts["static_pdf_pages"] += indexed_pages

    for path in sorted(context_dir.glob("*.docx")):
        text = _docx_text(path)
        if not text.strip():
            continue
        _add_document(
            documents,
            metadatas,
            ids,
            text=text,
            source=path.name,
            doc_type="static_report",
            key=path.stem,
            metadata={
                "title": path.stem.replace("_", " "),
                "file_type": "docx",
            },
        )
        counts["static_context_files"] += 1
        counts["static_docx_files"] += 1

    return documents, metadatas, ids, counts


def build_chroma_from_default_excel(flush_existing: bool = False) -> Dict[str, Any]:
    """Build the monolithic collection from the default Excel path (incremental-friendly)."""
    return build_chroma_complete(flush_existing=flush_existing)


def build_rag_corpus_complete(flush_existing: bool = True) -> Dict[str, Any]:
    """Alias for a full corpus rebuild; defaults to flushing the existing collection."""
    return build_chroma_complete(flush_existing=flush_existing)


def build_chroma_complete(
    flush_existing: bool = True,
    excel_path: Path = EXCEL_SOURCE,
    playbooks_dir: Path = PLAYBOOKS_DIR,
    static_context_dir: Path = STATIC_CONTEXT_DIR,
) -> Dict[str, Any]:
    """
    Build the monolithic `electronics_supply_chain_knowledge` ChromaDB collection.

    Pipeline: Excel + playbooks + static files → chunk → bi-encoder embed → upsert.
    ChromaDB calls `get_embedding_model()` on each upsert batch; the bi-encoder
    converts chunk text to 384-dim vectors stored in the HNSW cosine index.

    Args:
        flush_existing: Delete the collection before rebuild when True.
        excel_path: Path to the Lite Master workbook.
        playbooks_dir: Directory of mitigation `.txt` playbooks.
        static_context_dir: Directory of supplemental PDF/DOCX reports.

    Returns:
        Summary dict with chunk counts, source breakdown, and embedding model used.
    """
    if flush_existing:
        # Drop the collection via the API rather than deleting the directory.
        # shutil.rmtree on a live PersistentClient causes WinError 32 on Windows
        # because the client holds an exclusive lock on chroma.sqlite3.
        try:
            get_chroma_client().delete_collection(name=DEFAULT_COLLECTION_NAME)
        except Exception:
            pass  # collection may not exist yet — that's fine

    workbook_docs, workbook_meta, workbook_ids, counts = _workbook_documents(
        excel_path
    )
    playbook_docs, playbook_meta, playbook_ids, playbook_count = (
        _playbook_documents(playbooks_dir)
    )
    context_docs, context_meta, context_ids, context_counts = (
        _static_context_documents(static_context_dir)
    )

    documents = workbook_docs + playbook_docs + context_docs
    metadatas = workbook_meta + playbook_meta + context_meta
    ids = workbook_ids + playbook_ids + context_ids
    if not documents:
        raise ValueError("No electronics knowledge documents were generated")

    # Guard against duplicate IDs (Chroma upsert rejects duplicates in one batch)
    seen_ids: set[str] = set()
    dedup_docs: List[str] = []
    dedup_meta: List[Dict[str, Any]] = []
    dedup_ids: List[str] = []
    for doc, meta, doc_id in zip(documents, metadatas, ids):
        unique_id = doc_id
        suffix = 0
        while unique_id in seen_ids:
            suffix += 1
            unique_id = hashlib.sha256(f"{doc_id}|{suffix}".encode()).hexdigest()
        seen_ids.add(unique_id)
        dedup_docs.append(doc)
        dedup_meta.append(meta)
        dedup_ids.append(unique_id)
    documents, metadatas, ids = dedup_docs, dedup_meta, dedup_ids

    client = get_chroma_client()
    collection = client.get_or_create_collection(
        name=DEFAULT_COLLECTION_NAME,
        embedding_function=get_embedding_model(),
        metadata={
            "hnsw:space": "cosine",
            "domain": "electronics_semiconductors",
            "dataset_owner": "Varun",
        },
    )

    batch_size = 100
    for start in range(0, len(documents), batch_size):
        collection.upsert(
            ids=ids[start : start + batch_size],
            documents=documents[start : start + batch_size],
            metadatas=metadatas[start : start + batch_size],
        )

    return {
        "collection": DEFAULT_COLLECTION_NAME,
        "chunks": collection.count(),
        "source_documents": (
            sum(counts.values())
            + playbook_count
            + context_counts["static_context_files"]
        ),
        "playbooks": playbook_count,
        **counts,
        **context_counts,
        "domain": "electronics_semiconductors",
        "dataset_owner": "Varun",
        "beauty_products_included": False,
        "embedding_model": resolve_embedding_model_name(),
    }


def query_chroma_rag(
    query: str,
    n_results: int = 5,
    collection_name: str = DEFAULT_COLLECTION_NAME,
    where: Dict[str, Any] | None = None,
) -> List[Dict[str, Any]]:
    """
    Stage-1 bi-encoder retrieval against the monolithic collection.

    Embeds `query` with the same bi-encoder used at ingest, then returns the
    top-N chunks by ascending cosine distance (lower = more similar). Does not
    run cross-encoder reranking; agents needing higher precision should use
    `src.rag.retriever.retrieve_and_rerank` on named collections instead.

    Args:
        query: Natural-language search string.
        n_results: Maximum hits to return (capped at collection size).
        collection_name: Chroma collection name (default monolithic store).
        where: Optional Chroma metadata filter.

    Returns:
        List of dicts with keys `text`, `metadata`, and `distance`.
    """
    client = get_chroma_client()
    try:
        collection = client.get_collection(
            name=collection_name,
            embedding_function=get_embedding_model(),
        )
    except Exception:
        return []

    count = collection.count()
    if count == 0:
        return []
    results = collection.query(
        query_texts=[query],
        n_results=min(n_results, count),
        where=where,
        include=["documents", "metadatas", "distances"],
    )

    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]
    distances = results.get("distances", [[]])[0]
    return [
        {"text": doc, "metadata": metadata, "distance": distance}
        for doc, metadata, distance in zip(documents, metadatas, distances)
    ]


# ── Day 8 Screen 6 read helpers ─────────────────────────────────────────────

RAGAS_SCORES_PATH = PROJECT_ROOT / "evaluation" / "ragas" / "ragas_scores_full.json"
RAGAS_GOLD_DATASET_PATH = PROJECT_ROOT / "evaluation" / "ragas" / "test_dataset.json"

_RAGAS_METRIC_THRESHOLDS = {
    "Faithfulness": 0.75,
    "Answer Relevance": 0.80,
    "Context Precision": 0.70,
    "Context Recall": 0.75,
}

_RAGAS_METRIC_KEY_MAP = {
    "faithfulness": "Faithfulness",
    "answer_relevancy": "Answer Relevance",
    "context_precision": "Context Precision",
    "context_recall": "Context Recall",
}

_NAMED_COLLECTIONS = (
    "historical_precedents",
    "export_control_corpus",
    "india_sourcing_corpus",
)


def fetch_corpus_health() -> List[Dict[str, Any]]:
    """Live collection.count() for all 3 named ChromaDB collections (Screen 6).

    Returns list of {name, docs, real, synth, last_ingested_at} dicts.
    real/synth split inferred from [SOURCE: SYNTHESIZED] tag in chunk text."""
    from src.rag.collections import COLLECTION_NAMES

    client = get_chroma_client()
    embed_fn = get_embedding_model()
    results: List[Dict[str, Any]] = []

    for cname in COLLECTION_NAMES.values():
        try:
            col = client.get_collection(name=cname, embedding_function=embed_fn)
        except Exception:
            results.append(
                {
                    "name": cname,
                    "docs": 0,
                    "real": 0,
                    "synth": 0,
                    "last_ingested_at": "not ingested",
                }
            )
            continue

        total = col.count()
        synth = 0
        if total > 0:
            try:
                batch = col.get(include=["documents"])
                docs = batch.get("documents") or []
                synth = sum(
                    1 for doc in docs if doc and "[SOURCE: SYNTHESIZED]" in doc
                )
            except Exception:
                synth = 0
        real = max(0, total - synth)
        meta = col.metadata or {}
        chroma_mtime = CHROMA_DIR.stat().st_mtime if CHROMA_DIR.exists() else None
        if chroma_mtime:
            from datetime import datetime, timezone
            ingested = datetime.fromtimestamp(chroma_mtime, tz=timezone.utc).strftime(
                "%Y-%m-%d %H:%M UTC"
            )
        else:
            ingested = meta.get("last_ingested_at") or "—"
        results.append(
            {
                "name": cname,
                "docs": total,
                "real": real,
                "synth": synth,
                "last_ingested_at": str(ingested),
            }
        )
    return results


def fetch_ragas_scorecard() -> List[Dict[str, Any]]:
    """Read persisted RAGAS Phase 3 full-mode scores (Screen 6 scorecard).

    Source: evaluation/ragas/ragas_scores_full.json overall metrics.
    Returns empty list when the evaluation file is absent (scope cut)."""
    if not RAGAS_SCORES_PATH.exists():
        return []
    try:
        payload = json.loads(RAGAS_SCORES_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []

    overall = payload.get("overall") or {}
    run_at = payload.get("run_at_utc") or ""
    tiles: List[Dict[str, Any]] = []
    for key, label in _RAGAS_METRIC_KEY_MAP.items():
        score = overall.get(key)
        if score is None:
            continue
        threshold = _RAGAS_METRIC_THRESHOLDS[label]
        tiles.append(
            {
                "metric": label,
                "score": round(float(score), 4),
                "threshold": threshold,
                "passed": float(score) >= threshold,
            }
        )
    if run_at and tiles:
        pass  # run_at retained in JSON source only — RagasScore schema unchanged
    return tiles


def fetch_gold_dataset(limit: int = 50) -> List[Dict[str, Any]]:
    """Read chunk-grounded gold QA rows from RAGAS Phase 1 output (Screen 6).

    Source: evaluation/ragas/test_dataset.json test_cases, with match inferred
    from ragas_scores_full.json per_case faithfulness when available."""
    if not RAGAS_GOLD_DATASET_PATH.exists():
        return []
    try:
        dataset = json.loads(RAGAS_GOLD_DATASET_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []

    per_case_match: Dict[str, bool] = {}
    if RAGAS_SCORES_PATH.exists():
        try:
            scores = json.loads(RAGAS_SCORES_PATH.read_text(encoding="utf-8"))
            for case in scores.get("per_case") or []:
                q = case.get("question")
                if q:
                    per_case_match[q] = float(case.get("faithfulness") or 0) >= 0.75
        except (OSError, json.JSONDecodeError):
            pass

    rows: List[Dict[str, Any]] = []
    for case in (dataset.get("test_cases") or [])[:limit]:
        question = case.get("question") or ""
        rows.append(
            {
                "question": question,
                "ground_truth": case.get("ground_truth") or "",
                "match": per_case_match.get(question, True),
                "source_collection": case.get("source_collection"),
                "source_chunk_id": case.get("source_chunk_id"),
                "query_style": case.get("query_style") or "natural_question",
            }
        )
    return rows
